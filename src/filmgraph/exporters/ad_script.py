"""
FilmGraph → AD (Audiodeskription) .docx exporter.

Writes a German-style audio description script. Each shot with a visual
description becomes a timecoded block::

    HH:MM:SS
    * <description>          (if marked as a scene change in custom)
    <description>
    "dialogue in quotes"

Usage::

    from filmgraph.exporters.ad_script import filmgraph_to_ad
    filmgraph_to_ad(fg, "lost_killers_ad.docx")

CLI::

    python -m filmgraph.exporters.ad_script movie.filmgraph.json -o ad.docx
"""
from __future__ import annotations

from pathlib import Path

from filmgraph.schema import FilmGraph


def _sec_to_hms(seconds: float) -> str:
    if seconds < 0:
        seconds = 0.0
    total = int(round(seconds))
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def filmgraph_to_ad(fg: FilmGraph, output_path: str) -> str:
    """Write a FilmGraph as an AD script .docx file."""
    from docx import Document

    doc = Document()

    doc.add_paragraph(f"Hörfilmfassung / AD Script: {fg.meta.title}")
    if fg.meta.original_title and fg.meta.original_title != fg.meta.title:
        doc.add_paragraph(f"Projekt: {fg.meta.original_title}")
    if fg.meta.duration:
        doc.add_paragraph(f"Dauer: {_sec_to_hms(fg.meta.duration)}")
    doc.add_paragraph("")

    for scene in fg.scenes:
        for shot in scene.shots:
            description = shot.visual.description or ""
            dialogue_texts = [dl.text for dl in shot.audio.dialogue]

            if not description and not dialogue_texts:
                continue

            doc.add_paragraph(_sec_to_hms(shot.start_time))

            if description:
                is_scene_change = bool(shot.custom.get("scene_change"))
                prefix = "* " if is_scene_change else ""
                doc.add_paragraph(f"{prefix}{description}")

            for text in dialogue_texts:
                # German quote style: „text"
                doc.add_paragraph(f"\u201e{text}\u201c")

            doc.add_paragraph("")

    doc.save(output_path)
    return output_path


# ─── CLI ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="Convert FilmGraph to AD script .docx")
    ap.add_argument("input", help="Path to .filmgraph.json file")
    ap.add_argument("-o", "--output", default=None, help="Output .docx path")
    args = ap.parse_args()

    fg = FilmGraph.from_json(Path(args.input).read_text(encoding="utf-8"))
    out = args.output or str(Path(args.input).with_suffix(".ad.docx"))
    filmgraph_to_ad(fg, out)
    total_shots = sum(s.shot_count for s in fg.scenes)
    print(f"{total_shots} AD blocks -> {out}")
