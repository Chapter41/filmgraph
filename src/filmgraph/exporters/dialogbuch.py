"""
FilmGraph → Dialogbuch .docx exporter.

Writes a German-style dubbing script (Dialogbuch) with alternating
timestamp and speaker-prefixed dialogue lines::

    00:00:30 - 00:00:35
    Frieda Meisner: Nun sind wir wieder hier in Liebitz.

Only shots that contain at least one DialogueLine are emitted. Speakers
are resolved to their display name via the character entity list.

Usage::

    from filmgraph.exporters.dialogbuch import filmgraph_to_dialogbuch
    filmgraph_to_dialogbuch(fg, "episode_dialogbuch.docx")

CLI::

    python -m filmgraph.exporters.dialogbuch movie.filmgraph.json -o dialogbuch.docx
"""
from __future__ import annotations

from pathlib import Path

from filmgraph.schema import FilmGraph


def _sec_to_hms(seconds: float) -> str:
    if seconds < 0:
        seconds = 0.0
    total_ms = int(round(seconds * 1000))
    h = total_ms // 3_600_000
    m = (total_ms % 3_600_000) // 60_000
    s = (total_ms % 60_000) // 1000
    return f"{h:02d}:{m:02d}:{s:02d}"


def filmgraph_to_dialogbuch(fg: FilmGraph, output_path: str) -> str:
    """Write a FilmGraph as a Dialogbuch .docx file."""
    from docx import Document

    doc = Document()

    name_lookup = {c.id: c.name for c in fg.entities.characters}

    doc.add_heading(fg.meta.title, level=1)
    if fg.meta.language:
        doc.add_paragraph(f"Sprache: {fg.meta.language}")
    doc.add_paragraph("")

    for scene in fg.scenes:
        for shot in scene.shots:
            if not shot.audio.dialogue:
                continue
            start = _sec_to_hms(shot.start_time)
            end = _sec_to_hms(shot.end_time)
            doc.add_paragraph(f"{start} - {end}")

            for dl in shot.audio.dialogue:
                speaker_name = (
                    name_lookup.get(dl.speaker, dl.speaker.replace("-", " ").title())
                    if dl.speaker else "Unknown"
                )
                doc.add_paragraph(f"{speaker_name}: {dl.text}")

            doc.add_paragraph("")

    doc.save(output_path)
    return output_path


# ─── CLI ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="Convert FilmGraph to Dialogbuch .docx")
    ap.add_argument("input", help="Path to .filmgraph.json file")
    ap.add_argument("-o", "--output", default=None)
    args = ap.parse_args()

    fg = FilmGraph.from_json(Path(args.input).read_text(encoding="utf-8"))
    out = args.output or str(Path(args.input).with_suffix(".dialogbuch.docx"))
    filmgraph_to_dialogbuch(fg, out)
    blocks = sum(
        1 for sc in fg.scenes for sh in sc.shots if sh.audio.dialogue
    )
    print(f"{blocks} dialogue blocks -> {out}")
