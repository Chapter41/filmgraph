"""
FilmGraph → CCSL .docx exporter.

Writes a Combined Continuity & Spotting List (CCSL) .docx file in the
shot-number layout:

    Shot# | Timing | Description | Dialogue

This is the simpler of the two CCSL layouts supported by the importer
(the other is the multi-reel banner format used on features). Most
downstream tools read both, and the flat layout loses no information
when there is only a single reel.

Usage::

    from filmgraph.exporters.ccsl import filmgraph_to_ccsl
    filmgraph_to_ccsl(fg, "movie_ccsl.docx")

CLI::

    python -m filmgraph.exporters.ccsl movie.filmgraph.json -o movie_ccsl.docx
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from filmgraph.schema import FilmGraph, Shot


def _sec_to_tc(seconds: float, fps: float = 24.0) -> str:
    """Seconds → HH:MM:SS:FF SMPTE timecode."""
    if seconds < 0:
        seconds = 0.0
    total_frames = int(round(seconds * fps))
    frames_per_hour = int(round(3600 * fps))
    frames_per_minute = int(round(60 * fps))
    frames_per_second = int(round(fps))

    h = total_frames // frames_per_hour
    rem = total_frames % frames_per_hour
    m = rem // frames_per_minute
    rem %= frames_per_minute
    s = rem // frames_per_second
    f = rem % frames_per_second
    return f"{h:02d}:{m:02d}:{s:02d}:{f:02d}"


def _format_cinematography(shot: Shot) -> str:
    """Build the CCSL 'Description' cell from cinematography + visual layers."""
    cine = shot.cinematography
    pieces: list[str] = []

    if cine.shot_type:
        # Map enum value (e.g. "over-the-shoulder") to CCSL abbreviation
        type_abbr = {
            "over-the-shoulder": "OTS",
            "point-of-view": "POV",
        }.get(cine.shot_type.value)
        if type_abbr:
            pieces.append(type_abbr)

    if cine.shot_size:
        pieces.append(cine.shot_size.value)

    if cine.camera_movement:
        mv = cine.camera_movement.value
        if mv and mv != "static":
            pieces.append(mv.upper())

    if cine.camera_angle:
        ang = cine.camera_angle.value
        if ang and ang != "eye-level":
            angle_abbr = {
                "high-angle": "HIGH ANGLE",
                "low-angle": "LOW ANGLE",
                "dutch-angle": "DUTCH",
                "birds-eye": "BIRDS-EYE",
            }.get(ang, ang.upper())
            pieces.append(angle_abbr)

    if shot.visual.description:
        pieces.append(shot.visual.description)

    return " — ".join(pieces) if pieces else ""


def _format_dialogue(shot: Shot, fg: FilmGraph) -> str:
    """Build the CCSL 'Dialogue' cell as NAME: text lines."""
    lines: list[str] = []
    name_lookup = {c.id: c.name for c in fg.entities.characters}
    for dl in shot.audio.dialogue:
        if dl.speaker:
            name = name_lookup.get(dl.speaker, dl.speaker.replace("-", " ").title())
            lines.append(f"{name.upper()}: {dl.text}")
        else:
            lines.append(dl.text)
    return "\n".join(lines)


def filmgraph_to_ccsl(
    fg: FilmGraph,
    output_path: str,
    *,
    fps: Optional[float] = None,
) -> str:
    """Write a FilmGraph as a CCSL .docx file.

    Args:
        fg: The FilmGraph to export.
        output_path: Destination .docx path.
        fps: Frame rate for SMPTE conversion (defaults to fg.meta.frame_rate or 24).

    Returns the output path.
    """
    from docx import Document

    f = fps or fg.meta.frame_rate or 24.0
    doc = Document()

    doc.add_heading(fg.meta.title, level=1)
    if fg.meta.original_title and fg.meta.original_title != fg.meta.title:
        doc.add_paragraph(fg.meta.original_title)
    doc.add_paragraph(f"Combined Continuity & Spotting List — {fps or f:g} fps")

    shots = [sh for sc in fg.scenes for sh in sc.shots]

    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    header_cells = table.rows[0].cells
    header_cells[0].text = "Shot"
    header_cells[1].text = "Timing"
    header_cells[2].text = "Description"
    header_cells[3].text = "Dialogue"

    for shot in shots:
        row = table.add_row().cells
        row[0].text = str(shot.order)
        row[1].text = shot.timecode_in or _sec_to_tc(shot.start_time, f)
        row[2].text = _format_cinematography(shot)
        row[3].text = _format_dialogue(shot, fg)

    doc.save(output_path)
    return output_path


# ─── CLI ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="Convert FilmGraph to CCSL .docx")
    ap.add_argument("input", help="Path to .filmgraph.json file")
    ap.add_argument("-o", "--output", default=None, help="Output .docx path")
    ap.add_argument("--fps", type=float, default=None)
    args = ap.parse_args()

    fg = FilmGraph.from_json(Path(args.input).read_text(encoding="utf-8"))
    out = args.output or str(Path(args.input).with_suffix(".ccsl.docx"))
    filmgraph_to_ccsl(fg, out, fps=args.fps)
    total_shots = sum(s.shot_count for s in fg.scenes)
    print(f"{total_shots} shots -> {out}")
