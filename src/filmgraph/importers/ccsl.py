"""
CCSL .docx → FilmGraph importer.

Parses Combined Continuity & Spotting Lists (CCSL) and converts them into
FilmGraph format.  Supports two CCSL layouts:

  - **Reel-banner format** (e.g. Fisherwoman / Dead of Winter):
    REEL X banners with 'Scene N' rows, multi-reel SMPTE timecodes.

  - **Shot-number format** (e.g. Bloody Tennis):
    Flat table with Shot# | Timing | Description | Dialogue.

Usage::

    from filmgraph.importers.ccsl import ccsl_to_filmgraph
    fg = ccsl_to_filmgraph("path/to/ccsl.docx")
    Path("movie.filmgraph.json").write_text(fg.to_json())

CLI::

    python -m filmgraph.importers.ccsl input/Fisherwoman.docx -o fisherwoman.filmgraph.json
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from filmgraph.schema import (
    Audio,
    Character,
    Cinematography,
    DataSource,
    DialogueLine,
    Entities,
    FilmGraph,
    FilmMeta,
    MusicCue,
    Scene,
    Shot,
    Visual,
)
from filmgraph.vocabulary import CameraAngle, CameraMovement, ShotSize, ShotType


# ─── Shot size / camera parsing ──────────────────────────────────────────

# Map raw CCSL abbreviations → ShotSize enum
_SHOT_SIZE_MAP: dict[str, ShotSize] = {
    "ECU": ShotSize.ECU,
    "BCU": ShotSize.BCU,
    "CU": ShotSize.CU,
    "MCU": ShotSize.MCU,
    "MS": ShotSize.MS,
    "MLS": ShotSize.MLS,      # Medium Long Shot → native MLS
    "MWS": ShotSize.MWS,
    "LS": ShotSize.WS,        # Long Shot → WS
    "WS": ShotSize.WS,
    "EWS": ShotSize.EWS,
    "ELS": ShotSize.EWS,      # Extreme Long Shot → EWS
    "AERIAL": ShotSize.AERIAL,
    "DRONE": ShotSize.AERIAL,
    "INSERT": ShotSize.INSERT,
}

_CAMERA_MOVEMENT_MAP: dict[str, CameraMovement] = {
    "STATIC": CameraMovement.STATIC,
    "PAN": CameraMovement.PAN,
    "TILT": CameraMovement.TILT,
    "DOLLY": CameraMovement.DOLLY,
    "DOLY": CameraMovement.DOLLY,  # common typo in CCSLs
    "TRUCK": CameraMovement.TRUCK,
    "TRACKING": CameraMovement.TRACKING,
    "TRACK": CameraMovement.TRACKING,
    "HANDHELD": CameraMovement.HANDHELD,
    "STEADICAM": CameraMovement.STEADICAM,
    "CRANE": CameraMovement.CRANE,
    "DRONE": CameraMovement.DRONE,
    "ZOOM": CameraMovement.ZOOM,
    "ZOOM IN": CameraMovement.ZOOM,
    "ZOOM OUT": CameraMovement.ZOOM,
    "ZOO IN": CameraMovement.ZOOM,  # typo seen in BT CCSL
}

_CAMERA_ANGLE_MAP: dict[str, CameraAngle] = {
    "LOW ANGLE": CameraAngle.LOW,
    "LA": CameraAngle.LOW,
    "HIGH ANGLE": CameraAngle.HIGH,
    "HA": CameraAngle.HIGH,
    "BIRD": CameraAngle.BIRD,
    "BIRDS-EYE": CameraAngle.BIRD,
    "OVERHEAD": CameraAngle.OVERHEAD,
    "DUTCH": CameraAngle.DUTCH,
    "WORM": CameraAngle.WORM,
    "GROUND": CameraAngle.GROUND,
}

_SHOT_TYPE_MAP: dict[str, ShotType] = {
    "OTS": ShotType.OTS,
    "OVER-THE-SHOULDER": ShotType.OTS,
    "POV": ShotType.POV,
}


def _parse_cinematography(
    description: str,
    raw_shot_size: Optional[str] = None,
    raw_camera_movement: Optional[str] = None,
    raw_modifier: Optional[str] = None,
) -> Cinematography:
    """Extract cinematography enums from CCSL description and optional parsed fields.

    Handles compound annotations like 'OTS MS' → shot_type=OTS, shot_size=MS.
    """
    shot_size: Optional[ShotSize] = None
    camera_movement: Optional[CameraMovement] = None
    camera_angle: Optional[CameraAngle] = None
    shot_type: Optional[ShotType] = None

    # Use pre-parsed fields if available (Bloody Tennis GT format)
    if raw_shot_size:
        upper_ss = raw_shot_size.upper().strip()
        # OTS is a shot_type, not a shot_size — handle compound
        if upper_ss in _SHOT_TYPE_MAP:
            shot_type = _SHOT_TYPE_MAP[upper_ss]
            # Don't set shot_size — let description parsing find the real one
        else:
            shot_size = _SHOT_SIZE_MAP.get(upper_ss)

    if raw_camera_movement:
        camera_movement = _CAMERA_MOVEMENT_MAP.get(raw_camera_movement.upper().strip())

    if raw_modifier:
        camera_angle = _CAMERA_ANGLE_MAP.get(raw_modifier.upper().strip())

    # Parse from description text
    desc_upper = description.upper().strip()
    # Tokenize for ordered parsing
    tokens = re.split(r'[\s,]+', desc_upper)

    # Shot type from description (if not already set)
    if shot_type is None:
        for token in tokens:
            if token in _SHOT_TYPE_MAP:
                shot_type = _SHOT_TYPE_MAP[token]
                break

    # Shot size from description — skip tokens that are shot types
    if shot_size is None:
        for token in tokens:
            if token in _SHOT_TYPE_MAP:
                continue  # skip OTS/POV, look for the actual size
            if token in _SHOT_SIZE_MAP:
                shot_size = _SHOT_SIZE_MAP[token]
                break

    if camera_movement is None:
        for token, cm in _CAMERA_MOVEMENT_MAP.items():
            if token in desc_upper:
                camera_movement = cm
                break

    if camera_angle is None:
        for token, ca in _CAMERA_ANGLE_MAP.items():
            if token in desc_upper:
                camera_angle = ca
                break

    return Cinematography(
        shot_size=shot_size,
        camera_movement=camera_movement,
        camera_angle=camera_angle,
        shot_type=shot_type,
    )


# ─── Timecode parsing ────────────────────────────────────────────────────

def _tc_to_sec(tc: str, fps: float = 24.0) -> Optional[float]:
    """HH:MM:SS:FF → float seconds. Returns None on parse failure."""
    if not tc:
        return None
    tc = tc.strip().split("\n")[0].split("\r")[0]
    parts = re.split(r"[:;.,]", tc)
    try:
        if len(parts) == 4:
            h, m, s, f = (int(p) for p in parts)
            return h * 3600 + m * 60 + s + f / fps
    except (ValueError, TypeError):
        pass
    return None


def _is_changeover(tc_abs: float, window: float = 10.0) -> bool:
    """True if this TC is within the last N seconds of its SMPTE hour."""
    return (tc_abs % 3600) >= (3600 - window)


# ─── Reel-format parser ──────────────────────────────────────────────────

def _parse_reel_format(
    table, fps: float, tc_offset: float
) -> tuple[list[Shot], list[DataSource]]:
    """Fisherwoman-style: REEL X banners, Scene N rows."""
    reel_groups: list[tuple[str, list]] = []
    current_label: Optional[str] = None
    current_rows: list = []

    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        c0 = cells[0] if cells else ""
        if "REEL" in c0 and not c0.startswith("Scene"):
            if current_label is not None:
                reel_groups.append((current_label, current_rows))
            current_label = c0
            current_rows = []
        elif c0.startswith("Scene") and len(cells) > 1:
            tc_abs = _tc_to_sec(cells[1], fps)
            if tc_abs is not None:
                # cells layout: [Scene#, TC_In, Visual, ..., Dialogue (col 7), ..., Music (col 13)]
                current_rows.append((
                    c0,
                    cells[1],
                    tc_abs,
                    cells[2].replace("\n", " ").strip() if len(cells) > 2 else "",
                    cells[7].replace("\n", " ").strip() if len(cells) > 7 else "",
                    cells[13].strip() if len(cells) > 13 else "",
                ))
    if current_label:
        reel_groups.append((current_label, current_rows))

    # Build reel offset table
    reel_lookup: dict[int, dict] = {}
    film_offset = 0.0
    for label, rows in reel_groups:
        reel_num = int(label.split()[-1])
        if reel_num == 0:
            continue
        content = [
            r for r in rows
            if int(r[2] / 3600) == reel_num and not _is_changeover(r[2])
        ]
        if not content:
            continue
        first_tc = content[0][2]
        last_tc = content[-1][2]
        reel_lookup[reel_num] = {"film_offset": film_offset, "first_tc": first_tc}
        film_offset += last_tc - first_tc

    def _film_time(tc_abs: float) -> Optional[float]:
        rn = int(tc_abs / 3600)
        ri = reel_lookup.get(rn)
        if ri is None:
            return None
        return ri["film_offset"] + (tc_abs - ri["first_tc"])

    shots: list[Shot] = []
    order = 1
    for label, rows in reel_groups:
        reel_num = int(label.split()[-1])
        if reel_num == 0:
            continue
        for scene_ref, tc_raw, tc_abs, visual, dialogue_raw, music in rows:
            if _is_changeover(tc_abs):
                continue
            film_t = _film_time(tc_abs)
            if film_t is None or film_t < 0:
                continue

            seek_time = max(0.0, film_t + tc_offset)
            cine = _parse_cinematography(visual)

            # Parse dialogue: "NAME: text" pattern
            audio = Audio(
                music_cues=[MusicCue(label=music)] if music else [],
            )
            if dialogue_raw:
                dm = re.match(r"^([A-Z][A-Z\s\-'\.]{1,30}):\s*(.*)", dialogue_raw)
                if dm:
                    audio.dialogue = [DialogueLine(
                        speaker=dm.group(1).strip().lower().replace(" ", "-"),
                        text=dm.group(2).strip(),
                        start_time=seek_time,
                        end_time=seek_time,  # single TC, no end time in CCSL
                    )]

            shots.append(Shot(
                id=f"ccsl-{order:04d}",
                order=order,
                start_time=seek_time,
                end_time=seek_time,  # backfilled later
                timecode_in=tc_raw,
                cinematography=cine,
                visual=Visual(description=visual),
                audio=audio,
                source="ccsl",
                custom={"ccsl_scene_ref": scene_ref, "reel": label},
            ))
            order += 1

    # Backfill end_time from next shot's start_time
    for i in range(len(shots) - 1):
        shots[i].end_time = shots[i + 1].start_time
    if shots:
        shots[-1].end_time = shots[-1].start_time + 5.0  # last shot: provisional

    sources = [DataSource(layer="shots", source="ccsl-reel-format")]
    return shots, sources


# ─── Shot-number format parser ───────────────────────────────────────────

def _parse_shot_number_format(
    table, fps: float, tc_offset: float
) -> tuple[list[Shot], list[DataSource]]:
    """Bloody Tennis-style: Shot# | Timing | Description | Dialogue."""
    timing_col = 1
    raw: list[tuple[str, str, float, str, str]] = []

    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if not cells:
            continue
        c0 = cells[0]
        # detect header
        if not c0.isdigit():
            for i, c in enumerate(cells):
                if "timing" in c.lower() or "timecode" in c.lower():
                    timing_col = i
                    break
            continue
        if len(cells) <= timing_col:
            continue
        tc = _tc_to_sec(cells[timing_col], fps)
        if tc is not None and tc > 0:
            visual = cells[2].replace("\n", " ").strip() if len(cells) > 2 else ""
            dialogue = cells[3].replace("\n", " ").strip() if len(cells) > 3 else ""
            raw.append((f"Shot {c0}", cells[timing_col], tc, visual, dialogue))

    if not raw:
        return [], []

    # Zero-reference: program start = min TC rounded to nearest hour
    first_tc = raw[0][2]
    program_start = int(first_tc / 3600) * 3600.0

    shots: list[Shot] = []
    for order_idx, (scene_ref, tc_raw, tc_abs, visual, dialogue_raw) in enumerate(raw, 1):
        if _is_changeover(tc_abs):
            continue
        film_t = tc_abs - program_start
        seek_time = max(0.0, film_t + tc_offset) if tc_offset else film_t

        cine = _parse_cinematography(visual)
        audio = Audio()
        if dialogue_raw:
            dm = re.match(r"^([A-Z][A-Z\s\-'\.]{1,30}):\s*(.*)", dialogue_raw)
            if dm:
                audio.dialogue = [DialogueLine(
                    speaker=dm.group(1).strip().lower().replace(" ", "-"),
                    text=dm.group(2).strip(),
                    start_time=seek_time,
                    end_time=seek_time,
                )]

        shots.append(Shot(
            id=f"ccsl-{order_idx:04d}",
            order=order_idx,
            start_time=seek_time,
            end_time=seek_time,  # backfilled
            timecode_in=tc_raw,
            cinematography=cine,
            visual=Visual(description=visual),
            audio=audio,
            source="ccsl",
            custom={"ccsl_scene_ref": scene_ref},
        ))

    # Backfill end times
    for i in range(len(shots) - 1):
        shots[i].end_time = shots[i + 1].start_time
    if shots:
        shots[-1].end_time = shots[-1].start_time + 5.0

    sources = [DataSource(layer="shots", source="ccsl-shot-number-format")]
    return shots, sources


# ─── Public API ───────────────────────────────────────────────────────────

def ccsl_to_filmgraph(
    docx_path: str,
    *,
    title: Optional[str] = None,
    fps: float = 24.0,
    tc_offset: float = 0.0,
) -> FilmGraph:
    """Parse a CCSL .docx and return a FilmGraph.

    Args:
        docx_path: Path to the CCSL .docx file.
        title: Film title (defaults to filename).
        fps: Frame rate for SMPTE timecode conversion.
        tc_offset: Seconds to add to film_time to get video seek_time.
                   E.g. -44.42 for Fisherwoman/Dead of Winter.
    """
    from docx import Document

    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError(f"No tables found in {docx_path}")
    table = doc.tables[0]

    # Auto-detect format
    has_reel = False
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        c0 = cells[0] if cells else ""
        if "REEL" in c0 and not c0.startswith("Scene"):
            has_reel = True
            break

    if has_reel:
        shots, sources = _parse_reel_format(table, fps, tc_offset)
    else:
        shots, sources = _parse_shot_number_format(table, fps, tc_offset)

    if not title:
        title = Path(docx_path).stem.replace("_", " ")

    # Extract unique characters from shot dialogue and characters_visible
    char_ids: set[str] = set()
    for shot in shots:
        for dl in shot.audio.dialogue:
            if dl.speaker:
                char_ids.add(dl.speaker)

    characters = [Character(id=cid, name=cid.replace("-", " ").title()) for cid in sorted(char_ids)]

    # Compute total duration
    duration = shots[-1].end_time if shots else 0.0

    return FilmGraph(
        meta=FilmMeta(
            title=title,
            duration=duration,
            frame_rate=fps,
            sources=sources,
        ),
        entities=Entities(characters=characters),
        # All CCSL shots in a single scene (CCSL doesn't define narrative scenes)
        scenes=[
            Scene(
                id=1,
                title="Full Film (CCSL)",
                summary=f"All {len(shots)} CCSL shots",
                start_time=shots[0].start_time if shots else 0.0,
                end_time=shots[-1].end_time if shots else 0.0,
                shots=shots,
            )
        ] if shots else [],
    )


# ─── CLI ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse
    import sys

    ap = argparse.ArgumentParser(description="Convert CCSL .docx to FilmGraph JSON")
    ap.add_argument("docx", help="Path to CCSL .docx file")
    ap.add_argument("-o", "--output", default=None, help="Output .filmgraph.json path")
    ap.add_argument("--title", default=None, help="Film title")
    ap.add_argument("--fps", type=float, default=24.0)
    ap.add_argument("--tc-offset", type=float, default=0.0)
    args = ap.parse_args()

    fg = ccsl_to_filmgraph(args.docx, title=args.title, fps=args.fps, tc_offset=args.tc_offset)

    out = args.output or Path(args.docx).with_suffix(".filmgraph.json").name
    Path(out).write_text(fg.to_json(), encoding="utf-8")
    total_shots = sum(s.shot_count for s in fg.scenes)
    print(f"✅ {total_shots} shots → {out}")
