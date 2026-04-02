"""
AD (Audiodeskription) .docx → FilmGraph importer.

Parses German audio description scripts (Hörfilmfassung) which have:
  - Single timecodes (HH:MM:SS) on their own line
  - Visual descriptions (prefixed with * for scene changes)
  - Dialogue snippets in quotes
  - ATMER markers (breathing pauses)

Usage::

    from filmgraph.importers.ad_script import ad_to_filmgraph
    fg = ad_to_filmgraph("lost_killers_ad.docx")

CLI::

    python -m filmgraph.importers.ad_script ad_script.docx -o output.filmgraph.json
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from filmgraph.schema import (
    Audio,
    Character,
    DataSource,
    DialogueLine,
    Entities,
    FilmGraph,
    FilmMeta,
    Scene,
    Shot,
    Visual,
)


# ─── Patterns ────────────────────────────────────────────────────────────

_TC_LINE = re.compile(r"^(\d{2}:\d{2}:\d{2})$")
_DIALOGUE = re.compile(r'^[„""](.+?)["""]$')  # German/English quotes
_SCENE_CHANGE = re.compile(r"^\*\s*(.*)")  # * marks scene changes


def _hms_to_sec(hms: str) -> float:
    h, m, s = hms.split(":")
    return int(h) * 3600 + int(m) * 60 + int(s)


# ─── Main parser ────────────────────────────────────────────────────────

def ad_to_filmgraph(
    docx_path: str,
    *,
    title: Optional[str] = None,
    language: str = "de",
) -> FilmGraph:
    """Parse an AD script .docx and return a FilmGraph.

    Each timecoded block becomes a shot.
    Visual descriptions go into shot.visual.description.
    Dialogue snippets go into shot.audio.dialogue.
    """
    from docx import Document

    doc = Document(docx_path)

    # Extract title from header
    header_lines: list[str] = []
    for p in doc.paragraphs[:10]:
        text = p.text.strip()
        if text:
            header_lines.append(text)

    if not title:
        # Second line is usually the film title
        for line in header_lines:
            if "Projekt" not in line and "Auftrag" not in line and line:
                title = line.replace("_", " ")
                break
        if not title:
            title = Path(docx_path).stem

    # Parse blocks: timecode → content lines → next timecode
    lines = [p.text.strip() for p in doc.paragraphs]

    blocks: list[tuple[float, list[str]]] = []
    current_tc: Optional[float] = None
    current_lines: list[str] = []

    for line in lines:
        tc_match = _TC_LINE.match(line)
        if tc_match:
            if current_tc is not None and current_lines:
                blocks.append((current_tc, current_lines))
            current_tc = _hms_to_sec(tc_match.group(1))
            current_lines = []
        elif line and current_tc is not None:
            current_lines.append(line)

    if current_tc is not None and current_lines:
        blocks.append((current_tc, current_lines))

    # Convert blocks to shots
    shots: list[Shot] = []
    char_names: set[str] = set()

    for idx, (tc, content_lines) in enumerate(blocks):
        # End time = next block's start or +5s for last block
        end_time = blocks[idx + 1][0] if idx + 1 < len(blocks) else tc + 5.0

        descriptions: list[str] = []
        dialogue_lines: list[DialogueLine] = []
        is_scene_change = False

        for cl in content_lines:
            # Skip ATMER markers
            if cl.strip().upper() in ("ATMER", "GERÄUSCH"):
                continue

            # Scene change marker
            sc_match = _SCENE_CHANGE.match(cl)
            if sc_match:
                is_scene_change = True
                desc = sc_match.group(1).strip()
                if desc:
                    descriptions.append(desc)
                continue

            # Dialogue in quotes
            dl_match = _DIALOGUE.match(cl)
            if dl_match:
                dialogue_lines.append(DialogueLine(
                    text=dl_match.group(1),
                    start_time=tc,
                    end_time=end_time,
                    language=language,
                ))
                continue

            # Everything else is a visual description
            descriptions.append(cl)

        if not descriptions and not dialogue_lines:
            continue

        visual_text = " ".join(descriptions) if descriptions else None

        shots.append(Shot(
            id=f"ad-{len(shots) + 1:04d}",
            order=len(shots) + 1,
            start_time=tc,
            end_time=end_time,
            visual=Visual(description=visual_text),
            audio=Audio(dialogue=dialogue_lines) if dialogue_lines else Audio(),
            source="audiodeskription",
            custom={"scene_change": True} if is_scene_change else {},
        ))

    # Duration
    duration = shots[-1].end_time if shots else 0.0

    return FilmGraph(
        meta=FilmMeta(
            title=title,
            duration=duration,
            language=language,
            sources=[DataSource(layer="visual", source="audiodeskription")],
        ),
        entities=Entities(),
        scenes=[
            Scene(
                id=1,
                title="Full Film (AD Script)",
                summary=f"All {len(shots)} AD narration blocks",
                start_time=shots[0].start_time if shots else 0.0,
                end_time=shots[-1].end_time if shots else 0.0,
                shots=shots,
            )
        ] if shots else [],
    )


# ─── CLI ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="Convert AD script .docx to FilmGraph JSON")
    ap.add_argument("docx", help="Path to AD script .docx file")
    ap.add_argument("-o", "--output", default=None, help="Output .filmgraph.json path")
    ap.add_argument("--title", default=None, help="Film title")
    ap.add_argument("--language", default="de")
    args = ap.parse_args()

    fg = ad_to_filmgraph(args.docx, title=args.title, language=args.language)

    out = args.output or Path(args.docx).with_suffix(".filmgraph.json").name
    Path(out).write_text(fg.to_json(), encoding="utf-8")
    total_shots = sum(s.shot_count for s in fg.scenes)
    scene_changes = sum(
        1 for sc in fg.scenes for sh in sc.shots if sh.custom.get("scene_change")
    )
    print(f"✅ {total_shots} AD blocks ({scene_changes} scene changes) → {out}")
