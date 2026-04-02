"""
Dialogbuch .docx → FilmGraph importer.

Parses German dubbing scripts (Dialogbücher) which have a simple format:
  - Header lines (title, date, season/episode)
  - Alternating: timestamp line + speaker: dialogue line

Example::

    00:00:30 - 00:00:35
    Frieda Meisner: Nun sind wir wieder hier in Liebitz.

Usage::

    from filmgraph.importers.dialogbuch import dialogbuch_to_filmgraph
    fg = dialogbuch_to_filmgraph("001_FRIEDA_Dialogbuch.docx")

CLI::

    python -m filmgraph.importers.dialogbuch 001_FRIEDA_Dialogbuch.docx -o frieda_ep01.filmgraph.json
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from filmgraph.schema import (
    Audio,
    Character,
    DataSource,
    DialogueLine,
    Entities,
    FilmGraph,
    FilmMeta,
    Scene,
    Shot,
    Visual,
)


# ─── Time parsing ────────────────────────────────────────────────────────

_TC_PATTERN = re.compile(
    r"(\d{1,2}:\d{2}:\d{2}(?:\.\d+)?)\s*[-–]\s*(\d{1,2}:\d{2}:\d{2}(?:\.\d+)?)"
)


def _hms_to_sec(hms: str) -> float:
    """HH:MM:SS or H:MM:SS → float seconds."""
    parts = hms.split(":")
    if len(parts) == 3:
        h, m, s = parts
        return int(h) * 3600 + int(m) * 60 + float(s)
    elif len(parts) == 2:
        m, s = parts
        return int(m) * 60 + float(s)
    return float(parts[0])


# ─── Speaker parsing ────────────────────────────────────────────────────

_SPEAKER_PATTERN = re.compile(
    r"^([A-ZÄÖÜ][A-Za-zÄÖÜäöüß\s\-\.]+?):\s+(.*)", re.DOTALL
)


def _to_char_id(name: str) -> str:
    """Slugify a character name."""
    slug = re.sub(r"[^a-z0-9äöüß]+", "-", name.lower()).strip("-")
    # Normalize German umlauts for ID
    slug = slug.replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
    return slug


# ─── Main parser ────────────────────────────────────────────────────────

def dialogbuch_to_filmgraph(
    docx_path: str,
    *,
    title: Optional[str] = None,
    language: str = "de",
) -> FilmGraph:
    """Parse a Dialogbuch .docx and return a FilmGraph.

    Args:
        docx_path: Path to the Dialogbuch .docx file.
        title: Title override (defaults to filename parsing).
        language: Language code (default 'de' for German dubbing scripts).
    """
    from docx import Document

    doc = Document(docx_path)

    # Extract metadata from header paragraphs
    header_lines: list[str] = []
    for p in doc.paragraphs[:10]:
        text = p.text.strip()
        if text:
            header_lines.append(text)

    # Try to parse title from header
    if not title:
        # First non-empty line is usually the show name
        title = header_lines[0] if header_lines else Path(docx_path).stem

    # Parse episode info from headers
    episode_info = {}
    for line in header_lines:
        if "Staffel" in line:
            m = re.search(r"Staffel\s*(\d+)", line)
            if m:
                episode_info["season"] = int(m.group(1))
        if "Folge" in line:
            m = re.search(r"Folge\s*(\d+)", line)
            if m:
                episode_info["episode"] = int(m.group(1))

    # Parse alternating timestamp + dialogue lines
    shots: list[Shot] = []
    char_names: set[str] = set()
    order = 1

    # Collect all paragraph texts
    lines = [p.text.strip() for p in doc.paragraphs]

    i = 0
    while i < len(lines):
        line = lines[i]

        # Try to match timestamp
        tc_match = _TC_PATTERN.match(line)
        if tc_match:
            start_time = _hms_to_sec(tc_match.group(1))
            end_time = _hms_to_sec(tc_match.group(2))

            # Collect dialogue lines following this timestamp
            # (could be multiple speakers in one time block)
            dialogue_lines: list[DialogueLine] = []
            i += 1
            while i < len(lines):
                next_line = lines[i]
                if not next_line:
                    i += 1
                    continue
                # If it's a new timestamp, stop
                if _TC_PATTERN.match(next_line):
                    break
                # Try to extract speaker: text
                sp_match = _SPEAKER_PATTERN.match(next_line)
                if sp_match:
                    speaker_name = sp_match.group(1).strip()
                    text = sp_match.group(2).strip()
                    char_id = _to_char_id(speaker_name)
                    char_names.add(speaker_name)
                    dialogue_lines.append(DialogueLine(
                        speaker=char_id,
                        text=text,
                        start_time=start_time,
                        end_time=end_time,
                        language=language,
                    ))
                else:
                    # Continuation of previous dialogue or stage direction
                    if dialogue_lines:
                        dialogue_lines[-1].text += " " + next_line
                    # else: skip non-dialogue line
                i += 1

            if dialogue_lines:
                shots.append(Shot(
                    id=f"db-{order:04d}",
                    order=order,
                    start_time=start_time,
                    end_time=end_time,
                    audio=Audio(dialogue=dialogue_lines),
                    source="dialogbuch",
                ))
                order += 1
        else:
            i += 1

    # Build character list
    characters = [
        Character(id=_to_char_id(name), name=name)
        for name in sorted(char_names)
    ]

    # Zero-reference: some Dialogbücher use absolute TC (e.g. 10:00:01 for ep 3)
    # Find program start = first timestamp rounded down to nearest hour
    if shots:
        first_time = shots[0].start_time
        program_start = int(first_time / 3600) * 3600.0
        if program_start > 0:
            for shot in shots:
                offset = program_start
                shot.start_time -= offset
                shot.end_time -= offset
                for dl in shot.audio.dialogue:
                    dl.start_time -= offset
                    dl.end_time -= offset

    # Duration from last shot
    duration = shots[-1].end_time if shots else 0.0

    # Build episode title
    full_title = title
    if "season" in episode_info or "episode" in episode_info:
        parts = []
        if "season" in episode_info:
            parts.append(f"S{episode_info['season']:02d}")
        if "episode" in episode_info:
            parts.append(f"E{episode_info['episode']:02d}")
        full_title = f"{title} {''.join(parts)}"

    return FilmGraph(
        meta=FilmMeta(
            title=full_title,
            duration=duration,
            language=language,
            sources=[DataSource(layer="dialogue", source="dialogbuch")],
        ),
        entities=Entities(characters=characters),
        scenes=[
            Scene(
                id=1,
                title=f"Full Episode (Dialogbuch)",
                summary=f"All {len(shots)} dialogue blocks",
                start_time=shots[0].start_time if shots else 0.0,
                end_time=shots[-1].end_time if shots else 0.0,
                shots=shots,
            )
        ] if shots else [],
    )


# ─── CLI ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="Convert Dialogbuch .docx to FilmGraph JSON")
    ap.add_argument("docx", help="Path to Dialogbuch .docx file")
    ap.add_argument("-o", "--output", default=None, help="Output .filmgraph.json path")
    ap.add_argument("--title", default=None, help="Show/film title")
    ap.add_argument("--language", default="de")
    args = ap.parse_args()

    fg = dialogbuch_to_filmgraph(args.docx, title=args.title, language=args.language)

    out = args.output or Path(args.docx).with_suffix(".filmgraph.json").name
    Path(out).write_text(fg.to_json(), encoding="utf-8")
    total_shots = sum(s.shot_count for s in fg.scenes)
    chars = len(fg.entities.characters)
    print(f"✅ {total_shots} dialogue blocks, {chars} characters → {out}")
